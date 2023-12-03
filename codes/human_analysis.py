from docx.shared import RGBColor
from docx import Document
import re
import glob

doc_lst = glob.glob("lda_words/*.docx")

# Define a function to check if a run is red
def is_color(run, color):
    return {
        "red": run.font.color.rgb == RGBColor(255, 0, 0),
        "green": run.font.color.rgb == RGBColor(0, 255, 0) or run.font.color.rgb == RGBColor(0, 176, 80),
        "black": run.font.color.rgb == RGBColor(0, 0, 0) or run.font.color.rgb == None,
        "default": run.font.color.rgb == None
    }[color]

def save_doc(path_name, word_lst):
    doc = Document()
    doc.add_paragraph(' '.join(word_lst))
    doc.save(path_name)
    print("Saved " + path_name)

for doc_path in doc_lst:
    # Load a document
    doc = Document(doc_path)

    # Reinitialize the list for red words
    red_words, green_words, black_words = [], [], []

    # Iterate through the document to find red words
    for paragraph in doc.paragraphs:
        # Each paragraph stats with Row ##: ... Delete the Row ##: part
        row_num = paragraph.text.split(": ")[0]
        # paragraph.text = paragraph.text.split(": ")[1]
        red_words.append("\n\n" + row_num + ":")
        green_words.append("\n\n" + row_num + ":")
        black_words.append("\n\n")
        for run in paragraph.runs: 
            if is_color(run, "red"):
                # Adding space-separated words from the run
                red_words.extend(re.findall(r'\w+|:|-', run.text))
            elif is_color(run, "black"):
                black_words.extend(re.findall(r'\w+|:|-', run.text))
            else:
                green_words.extend(re.findall(r'\w+|:|-', run.text))

    # Save the new document with red words
    save_doc("./labels/" + doc_path.split("/")[1].replace(".docx", '') + "_negative.docx", red_words)
    save_doc("./labels/" + doc_path.split("/")[1].replace(".docx", '') + "_positive.docx", green_words)
    save_doc("./labels/" + doc_path.split("/")[1].replace(".docx", '') + "_neutral.docx", black_words)

def extract_paragraphs(docx_file):
    # Open the Word document
    doc = Document(docx_file)

    # Extract paragraphs
    paragraphs = {para.text.split(": ")[0]: para.text.split(": ")[1] for para in doc.paragraphs}

    return paragraphs